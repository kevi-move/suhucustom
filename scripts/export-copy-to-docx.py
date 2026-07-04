#!/usr/bin/env python3
"""Export public-site copy from suhucustom source into one .docx per page."""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Installing python-docx...", file=sys.stderr)
    import subprocess

    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
OUT_DIR = ROOT / "suhucustom-文案导出"

# --- navigation (mirrors src/lib/navigation.ts) ---
SERVICE_GROUPS = [
    {
        "titleEn": "Tops & Activewear",
        "items": [
            {"slug": "t-shirts", "nameEn": "T-shirts"},
            {"slug": "hoodies-sweatshirts", "nameEn": "Hoodies & Sweatshirts"},
            {"slug": "activewear-athleisure", "nameEn": "Activewear & Athleisure"},
            {"slug": "gym-sportswear", "nameEn": "Gym & Sportswear"},
        ],
    },
    {
        "titleEn": "Bottoms & Underwear",
        "items": [
            {"slug": "leggings", "nameEn": "Leggings"},
            {"slug": "jeans-denim", "nameEn": "Jeans & Denim"},
            {"slug": "underwear-bras", "nameEn": "Underwear & Bras"},
            {"slug": "swimwear", "nameEn": "Swimwear"},
        ],
    },
    {
        "titleEn": "Accessories & Headwear",
        "items": [
            {"slug": "hats-headwear", "nameEn": "Hats & Headwear"},
            {"slug": "socks", "nameEn": "Socks"},
            {"slug": "neck-gaiters", "nameEn": "Neck Gaiters"},
            {"slug": "leather-goods", "nameEn": "Leather Goods"},
        ],
    },
    {
        "titleEn": "Specialized & Home",
        "items": [
            {"slug": "uniforms", "nameEn": "Uniforms"},
            {"slug": "baby-kids-clothing", "nameEn": "Baby & Kids Clothing"},
            {"slug": "towels", "nameEn": "Towels"},
            {"slug": "cushion-covers", "nameEn": "Cushion Covers"},
        ],
    },
]

CUSTOMIZATION_ITEMS = [
    {"slug": "printing", "nameEn": "Printing"},
    {"slug": "embroidery", "nameEn": "Embroidery"},
    {"slug": "private-label", "nameEn": "Private Label"},
    {"slug": "tech-pack-design", "nameEn": "Tech Pack Design"},
]

SERVICE_COMPONENT_DIRS: dict[str, str] = {
    "t-shirts": "components/services/tshirts",
    "hoodies-sweatshirts": "components/services/hoodies",
    "activewear-athleisure": "components/services/activewear",
    "gym-sportswear": "components/services/gym",
    "leggings": "components/services/leggings",
    "jeans-denim": "components/services/jeans",
    "socks": "components/services/socks",
    "neck-gaiters": "components/services/neck-gaiters",
    "leather-goods": "components/services/leather-goods",
    "towels": "components/services/towels",
    "cushion-covers": "components/services/cushion-covers",
}

GENERIC_SERVICE_SLUGS = {
    "underwear-bras",
    "swimwear",
    "hats-headwear",
    "uniforms",
    "baby-kids-clothing",
}

SKIP_STRING_RE = re.compile(
    r"(^/|https?://|@/|\./|#|\.tsx?$|\.svg|\.png|\.jpg|YOUR_|wa\.me|"
    r"placeholder|object-|flex-|grid-|max-w-|px-|py-|sm:|lg:|"
    r"aria-|lucide|currentColor|viewBox|xmlns|className|href=|src=)",
    re.I,
)

TAILWIND_LIKE = re.compile(r"^[a-z0-9]+(?:-[a-z0-9]+)+$")


def find_service(slug: str) -> tuple[dict, dict] | None:
    for group in SERVICE_GROUPS:
        for item in group["items"]:
            if item["slug"] == slug:
                return item, group
    return None


def has_chinese(text: str) -> bool:
    return bool(re.search(r"[\u4e00-\u9fff]", text))


def decode_js_string(raw: str) -> str:
    s = raw[1:-1]
    s = s.replace("\\n", "\n").replace('\\"', '"').replace("\\'", "'")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def is_copy_candidate(s: str) -> bool:
    if not s or len(s) < 2:
        return False
    if has_chinese(s):
        return False
    if not re.search(r"[a-zA-Z]", s):
        return False
    if SKIP_STRING_RE.search(s):
        return False
    if TAILWIND_LIKE.match(s) and " " not in s:
        return False
    if len(s) > 800:
        return False
    if re.fullmatch(r"[\d\s\W]+", s):
        return False
    return True


def extract_from_source(text: str) -> list[str]:
    found: list[str] = []

    for m in re.finditer(r'"(?:[^"\\]|\\.)*"', text):
        s = decode_js_string(m.group(0))
        if is_copy_candidate(s):
            found.append(s)

    for m in re.finditer(r"'(?:[^'\\]|\\.)*'", text):
        s = decode_js_string(m.group(0))
        if is_copy_candidate(s):
            found.append(s)

    for m in re.finditer(r">([^<>{}]+)<", text):
        chunk = m.group(1).strip()
        chunk = re.sub(r"\s+", " ", chunk)
        if is_copy_candidate(chunk):
            found.append(chunk)

    for m in re.finditer(
        r"(?:title|desc|description|label|eyebrow|subtitle|paragraph\d*|intro|closing|"
        r"benefits|capabilities|primaryCtaText|secondaryCtaText|excerpt|placeholder|alt)"
        r'\s*:\s*"([^"]+)"',
        text,
        re.I,
    ):
        s = decode_js_string('"' + m.group(1) + '"')
        if is_copy_candidate(s):
            found.append(s)

    for m in re.finditer(
        r"(?:title|desc|description|label|eyebrow|subtitle|paragraph\d*|intro|closing|"
        r"benefits|capabilities|primaryCtaText|secondaryCtaText)\s*:\s*\n\s*\"([^\"]+)\"",
        text,
        re.I,
    ):
        s = m.group(1).strip()
        if is_copy_candidate(s):
            found.append(s)

    return dedupe_preserve_order(found)


def dedupe_preserve_order(items: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for item in items:
        key = item.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(item)
    return out


def read_files(paths: list[Path]) -> list[str]:
    all_text: list[str] = []
    for p in paths:
        if p.is_file():
            all_text.extend(extract_from_source(p.read_text(encoding="utf-8", errors="replace")))
    return dedupe_preserve_order(all_text)


def glob_tsx(dir_rel: str) -> list[Path]:
    d = SRC / dir_rel
    if not d.exists():
        return []
    return sorted(d.glob("**/*.tsx"))


def extract_about_us() -> list[str]:
    defaults = (SRC / "lib" / "aboutUsDefaults.ts").read_text(encoding="utf-8", errors="replace")
    strings = extract_from_source(defaults)
    about_components = glob_tsx("components/about")
    strings.extend(read_files(about_components))
    app_files = [
        SRC / "app" / "about-us" / "page.tsx",
        SRC / "app" / "about-us" / "AboutUsPageClient.tsx",
    ]
    strings.extend(read_files([p for p in app_files if p.exists()]))
    return dedupe_preserve_order(strings)


def generic_service_copy(slug: str) -> list[str]:
    found = find_service(slug)
    if not found:
        return []
    item, group = found
    name = item["nameEn"]
    return [
        "Services",
        name,
        group["titleEn"],
        name,
        group["titleEn"],
        f"Professional manufacturing for {name.lower()}. Custom specifications, quality materials, and reliable delivery for your brand or business.",
        "Request a Quote",
        "View Quality Process",
        "Image placeholder (you can replace with real product photos later)",
    ]


def customization_slug_copy(slug: str) -> list[str]:
    item = next((i for i in CUSTOMIZATION_ITEMS if i["slug"] == slug), None)
    if not item:
        return []
    name = item["nameEn"]
    return [
        "Home",
        "Customization",
        name,
        name,
        f"Our {name.lower()} services help you create unique, branded products that stand out in the market.",
        "Get Started",
    ]


def services_list_copy() -> list[str]:
    lines = [
        "Our Services",
        "Comprehensive apparel manufacturing for every category. Browse by category or find your product below.",
    ]
    for group in SERVICE_GROUPS:
        lines.append(group["titleEn"])
        for item in group["items"]:
            lines.append(item["nameEn"])
    return lines


def customization_list_copy() -> list[str]:
    lines = [
        "Customization Services",
        "From printing to tech packs, we offer full customization solutions for your apparel brand.",
    ]
    for item in CUSTOMIZATION_ITEMS:
        lines.append(item["nameEn"])
    return lines


def write_docx(filename: str, title: str, route: str, lines: list[str]) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title, level=0)
    p = doc.add_paragraph()
    p.add_run("Route: ").bold = True
    p.add_run(route)

    doc.add_paragraph()
    doc.add_heading("Page copy (English)", level=1)

    for line in lines:
        if "\n" in line:
            for part in line.split("\n"):
                if part.strip():
                    doc.add_paragraph(part.strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)

    path = OUT_DIR / filename
    doc.save(path)


# Page manifest: (index, chinese label, filename stem, route, copy lines)
def build_pages() -> list[tuple[int, str, str, str, list[str]]]:
    pages: list[tuple[int, str, str, str, list[str]]] = []
    idx = 1

    def add(label: str, route: str, lines: list[str], stem: str | None = None):
        nonlocal idx
        num = f"{idx:02d}"
        fname = f"{num}-{label}.docx" if stem is None else f"{num}-{stem}.docx"
        pages.append((idx, label, fname, route, lines))
        idx += 1

    home_files = [SRC / "app" / "page.tsx", *glob_tsx("components/home")]
    add("首页", "/", read_files(home_files), "首页")

    add("服务列表", "/services", services_list_copy(), "服务列表")

    for group in SERVICE_GROUPS:
        for item in group["items"]:
            slug = item["slug"]
            route = f"/services/{slug}"
            label = f"服务-{item['nameEn']}"
            if slug in SERVICE_COMPONENT_DIRS:
                rel = SERVICE_COMPONENT_DIRS[slug]
                lines = read_files(glob_tsx(rel))
            else:
                lines = generic_service_copy(slug)
            add(label, route, lines, f"服务-{slug}")

    add("定制列表", "/customization", customization_list_copy(), "定制列表")

    for item in CUSTOMIZATION_ITEMS:
        slug = item["slug"]
        add(
            f"定制-{item['nameEn']}",
            f"/customization/{slug}",
            customization_slug_copy(slug),
            f"定制-{slug}",
        )

    add("关于我们", "/about-us", extract_about_us(), "关于我们")

    blog_files = [
        SRC / "app" / "blog" / "page.tsx",
        *glob_tsx("components/blog"),
    ]
    blog_lines = read_files(blog_files)
    # Blog post detail shell (static UI only; article body is from CMS)
    blog_detail = read_files(
        [
            SRC / "app" / "blog" / "[slug]" / "page.tsx",
            SRC / "components" / "blog" / "BlogPostBody.tsx",
            SRC / "components" / "blog" / "BlogRelatedPosts.tsx",
        ]
    )
    add("博客列表", "/blog", blog_lines, "博客列表")
    add(
        "博客文章模板",
        "/blog/[slug]",
        blog_detail,
        "博客-文章模板",
    )

    add(
        "联系我们",
        "/contact-us",
        read_files([SRC / "app" / "contact-us" / "page.tsx"]),
        "联系我们",
    )

    add(
        "质量与生产",
        "/quality",
        read_files([SRC / "app" / "quality" / "page.tsx"]),
        "质量与生产",
    )

    add(
        "案例研究",
        "/company/case-studies",
        read_files([SRC / "app" / "company" / "case-studies" / "page.tsx"]),
        "案例研究",
    )

    search_files = [
        SRC / "app" / "search" / "page.tsx",
        SRC / "components" / "search" / "SearchForm.tsx",
        SRC / "components" / "search" / "SearchResultsList.tsx",
    ]
    add("搜索", "/search", read_files(search_files), "搜索")

    add(
        "隐私政策",
        "/privacy-policy",
        read_files([SRC / "app" / "privacy-policy" / "page.tsx"]),
        "隐私政策",
    )

    add(
        "条款与条件",
        "/terms-and-conditions",
        read_files([SRC / "app" / "terms-and-conditions" / "page.tsx"]),
        "条款与条件",
    )

    return pages


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    pages = build_pages()
    manifest = []

    for num, label, fname, route, lines in pages:
        if not lines:
            lines = ["(No extractable copy found in source for this page.)"]
        write_docx(fname, label, route, lines)
        manifest.append({"index": num, "label": label, "file": fname, "route": route, "lineCount": len(lines)})
        print(f"  {fname}  ({len(lines)} lines)")

    (OUT_DIR / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(f"\nDone: {len(pages)} documents -> {OUT_DIR}")


if __name__ == "__main__":
    main()
